import os
import random
from reportlab.pdfgen import canvas
from docx import Document

os.makedirs("sample_data", exist_ok=True)

# Sample phrases to populate resumes
SKILLS = [
    "Python, SQL, AWS, Java, C++",
    "Python, Django, Flask, HTML, CSS",
    "JavaScript, TypeScript, React, Node.js",
    "Java, Spring Boot, MySQL, Docker",
    "Ruby on Rails, PostgreSQL, Redis",
    "Go, Kubernetes, Python, AWS",
    "C#, .NET, Azure, SQL Server",
    "Python experience with Data Science, Pandas, NumPy",
]

EXPERIENCES = [
    "Software Engineer with 5 years of experience building scalable backends.",
    "Data Scientist specializing in machine learning and predictive modeling.",
    "Frontend Developer focused on modern UI/UX design and responsive web apps.",
    "DevOps Engineer with expertise in modern CI/CD pipelines.",
    "Full-stack developer who loves creating intuitive user experiences.",
]

def make_pdf(filename, content):
    c = canvas.Canvas(filename)
    c.drawString(100, 750, "RESUME")
    y = 700
    for line in content.split('\n'):
        c.drawString(100, y, line)
        y -= 20
    c.save()

def make_docx(filename, content):
    doc = Document()
    doc.add_heading('Resume', 0)
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(filename)

def make_txt(filename, content):
    with open(filename, 'w') as f:
        f.write(content)

names = ["john_doe", "jane_smith", "alex_davis", "maria_garcia", "david_wilson", "emma_brown"]

for i, name in enumerate(names):
    ext = random.choice([".pdf", ".txt", ".docx"])
    filepath = os.path.join("sample_data", f"resume_{name}{ext}")
    
    skill = random.choice(SKILLS)
    exp = random.choice(EXPERIENCES)
    
    content = f"Name: {name.replace('_', ' ').title()}\n\n"
    content += f"Summary:\n{exp}\n\n"
    content += f"Skills:\n{skill}\n\n"
    content += "Experience:\nWorked at XYZ Corp for 3 years.\nLed a team of 4 engineers."
    
    if ext == ".pdf":
        make_pdf(filepath, content)
    elif ext == ".docx":
        make_docx(filepath, content)
    else:
        make_txt(filepath, content)
        
print("Dummy data generated successfully in sample_data/")
